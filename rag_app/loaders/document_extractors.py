from io import BytesIO
from pathlib import Path
from typing import BinaryIO

import pandas as pd
from docx import Document
from pypdf import PdfReader

from rag_app import config


def validate_file_magic(file_bytes: bytes, filename: str) -> None:
    """验证文件魔数,防止伪装文件类型
    
    安全修复: 通过检查文件头部的magic bytes验证真实文件类型
    防止攻击者上传伪装成PDF/DOCX的恶意文件
    """
    suffix = Path(filename).suffix.lower().lstrip(".")
    
    # 文本文件不需要魔数验证
    if suffix in {"txt", "md", "csv"}:
        if len(file_bytes) < config.MIN_TEXT_FILE_SIZE:
            raise ValueError("文本文件内容为空。")
        return
    
    # 检查是否有对应的魔数配置
    if suffix not in config.FILE_MAGIC_SIGNATURES:
        # 没有配置魔数验证的文件类型,仅依赖扩展名验证
        return
    
    # 获取该文件类型的所有有效魔数
    valid_signatures = config.FILE_MAGIC_SIGNATURES[suffix]
    
    # 检查文件是否匹配任何一个有效魔数
    for signature in valid_signatures:
        if file_bytes[:len(signature)] == signature:
            return  # 验证通过
    
    # 所有魔数都不匹配
    raise ValueError(
        f"文件内容校验失败: {filename} 的实际格式与扩展名 .{suffix} 不匹配。"
        f"请确保文件未损坏且格式正确。"
    )


def extract_text_from_file(file_source: bytes | BinaryIO, filename: str) -> str:
    # 先读取文件字节用于验证
    file_bytes = _read_file_bytes(file_source)
    
    # 安全验证: 检查文件魔数
    validate_file_magic(file_bytes, filename)
    
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md", ".csv"}:
        return _decode_text_bytes(file_bytes)
    if suffix == ".pdf":
        return _extract_pdf_text(file_bytes)
    if suffix == ".docx":
        return _extract_docx_text(file_bytes)
    if suffix in {".xlsx", ".xls"}:
        return _extract_excel_text(file_bytes)
    if suffix == ".doc":
        raise ValueError("当前仅支持 .docx，老式 .doc 请先另存为 .docx 后再上传。")
    raise ValueError(f"暂不支持 {suffix or '无扩展名'} 格式。")


def _read_file_bytes(file_source: bytes | BinaryIO) -> bytes:
    if isinstance(file_source, bytes):
        return file_source

    current_position = file_source.tell() if file_source.seekable() else None
    if file_source.seekable():
        file_source.seek(0)
    file_bytes = file_source.read()
    if isinstance(file_bytes, str):
        file_bytes = file_bytes.encode("utf-8")
    if current_position is not None:
        file_source.seek(current_position)
    return bytes(file_bytes)


def _decode_text_bytes(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def _extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    parts = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        if page_text:
            parts.append(f"第{index}页\n{page_text}")
    return "\n\n".join(parts)


def _extract_docx_text(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"表格{table_index}\n" + "\n".join(rows))

    return "\n\n".join(parts)


def _extract_excel_text(file_bytes: bytes) -> str:
    sheets = pd.read_excel(BytesIO(file_bytes), sheet_name=None, dtype=str)
    parts = []

    for sheet_name, dataframe in sheets.items():
        dataframe = dataframe.fillna("")
        if dataframe.empty and not list(dataframe.columns):
            continue
        sheet_text = dataframe.to_csv(index=False).replace("\r\n", "\n").replace("\r", "\n").strip()
        if sheet_text:
            parts.append(f"工作表: {sheet_name}\n{sheet_text}")

    return "\n\n".join(parts)
